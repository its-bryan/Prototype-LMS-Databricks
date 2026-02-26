"""
Generate the MMR Flow Optimisation PRD as a Word document,
matching the example PRD template format.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:color"), val.get("color", "000000"))
            el.set(qn("w:space"), "0")
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading_styled(doc, text, level=1):
    """Add a heading with dark teal colour matching the example PRD."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)  # dark teal
    return h

def add_body(doc, text):
    """Add a normal paragraph."""
    return doc.add_paragraph(text)

def add_bold_body(doc, parts):
    """Add a paragraph with mixed bold/normal runs.
    parts is a list of (text, is_bold) tuples.
    """
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
    return p

def add_bullet(doc, parts, level=0):
    """Add a bullet point with mixed bold/normal runs."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
    return p

def add_numbered(doc, parts):
    """Add a numbered list item with mixed bold/normal runs."""
    p = doc.add_paragraph(style="List Number")
    p.clear()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    border_def = {"val": "single", "sz": "6", "color": "1B3A4B"}
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, "E8EDF0")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Support bold first column
            run = p.add_run(val)
            if c_idx == 0:
                run.bold = True
            run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ── Document Setup ───────────────────────────────────────────────────────

doc = Document()

# Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


# ── Title Page ───────────────────────────────────────────────────────────

p = doc.add_paragraph()
run = p.add_run("Product Requirements Document (PRD)")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
run = p.add_run("MMR Flow Optimisation — Insurance Replacement Conversion")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph("")  # spacer

# Table of Contents
h = doc.add_heading("Table of Contents", level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

toc_items = [
    "1. Context & Background",
    "2. Business Objectives",
    "3. Problem Statement",
    "4. Intended Behaviour",
    "5. Current User Experience & Flows (As-Is)",
    "6. Target User Experience & Flows (To-Be)",
    "7. MMR Booking Time Rules & Constraints",
    "8. Functional Requirements",
    "9. Non-Functional Requirements",
    "10. Success Metrics & Tracking Plan",
    "11. Out of Scope",
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph("")  # spacer

# Stakeholders
p = doc.add_paragraph(
    "The requirements outlined in this PRD have been formulated based on inputs "
    "gathered during multiple discovery and alignment sessions with key business "
    "stakeholders, including:"
)

stakeholders = [
    "Daniel Sia – Data & Analytics Lead, Insurance Replacement",
    "Randy – Operations, Insurance Replacement",
    "Anupam – Technical Product, MMR Platform",
    "Zac – Insurance Replacement Business",
    "Kalyan – Insurance Replacement Business",
]
for i, s in enumerate(stakeholders, 1):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    run = p.add_run(s)

doc.add_page_break()


# ── 1. Context & Background ─────────────────────────────────────────────

add_heading_styled(doc, "1. Context & Background")

add_body(doc, (
    "When a lead comes through Hertz's insurance partners, the customer automatically "
    "receives an SMS with a link to the Manage My Reservation (MMR) web application. "
    "The MMR allows customers to review their pre-populated reservation details and "
    "confirm their rental booking through a guided two-page flow."
))

add_bold_body(doc, [
    ("Customers who complete the MMR flow have a ", False),
    ("significantly higher conversion rate", True),
    (" than those who do not. ", False),
    ("Therefore, increasing MMR engagement and completion directly lifts overall "
     "insurance replacement conversion.", False),
])

add_body(doc, "The current MMR funnel performance is as follows:")

add_bullet(doc, [("SMS sent → 32% click-through rate", False)])
add_bullet(doc, [("Page 1 (My Details & Damaged Vehicle) → 87% completion", False)])
add_bullet(doc, [("Page 2 (Start a Reservation) → 66% completion", False)])

add_bold_body(doc, [
    ("The ", False),
    ("34% drop-off on Page 2", True),
    (" represents a significant loss of engaged customers who have already demonstrated "
     "intent by clicking the link and completing Page 1. This PRD defines the changes "
     "required to reduce Page 2 abandonment and increase SMS click-through, thereby "
     "lifting overall conversion.", False),
])


# ── 2. Business Objectives ──────────────────────────────────────────────

add_heading_styled(doc, "2. Business Objectives")

add_bullet(doc, [("Increase ", False), ("MMR SMS click-through rate", True),
                  (" from 32% toward 40–50%", False)])
add_bullet(doc, [("Increase ", False), ("MMR Page 2 completion rate", True),
                  (" from 66% toward 82–89%", False)])
add_bullet(doc, [("Reduce form abandonment", True),
                  (" by accommodating all customer scenarios, not just the body-shop-ready happy path", False)])
add_bullet(doc, [("Capture same-day rental demand", True),
                  (" (37% of leads) that is currently blocked by booking guardrails", False)])
add_bullet(doc, [("Secure early customer commitment", True),
                  (" to Hertz to reduce competitor flipping at the body shop", False)])

add_body(doc, "These objectives directly support:")

add_bullet(doc, [("Incremental conversion uplift of ", False),
                  ("+1.24 to +2.29 percentage points", True),
                  (" (MMR Page 1 + Page 2 combined)", False)])
add_bullet(doc, [("Estimated ", False), ("$1.2M–$2.3M EBITDA uplift", True)])
add_bullet(doc, [("Estimated ", False), ("$3.3M–$6.2M Revenue uplift", True)])
add_bullet(doc, [("Based on $1M EBITDA and $2.7M Revenue per 1% conversion improvement", False)])


# ── 3. Problem Statement ────────────────────────────────────────────────

add_heading_styled(doc, "3. Problem Statement")

p = doc.add_paragraph()
run = p.add_run("Page 2 Drop-Off")
run.bold = True

add_body(doc, (
    "87% of customers who click the MMR link complete Page 1, but only 66% complete "
    "Page 2 — a 34% drop-off rate. Page 2 has friction including confusing fields "
    "and an inflexible flow that only supports one customer journey."
))

p = doc.add_paragraph()
run = p.add_run("The MMR only supports customers who already have a body shop reservation.")
run.bold = True

add_body(doc, "Customers who fall outside this single happy path cannot complete the MMR flow:")

add_bullet(doc, [("Do not have a body shop reservation yet", False)])
add_bullet(doc, [("Do not know who their body shop will be", False)])
add_bullet(doc, [("Need to pick up today (same-day)", False)])
add_bullet(doc, [("Want to be picked up from somewhere other than the body shop "
                   "(e.g., home, work, other location)", False)])
add_bullet(doc, [("Do not need pickup at all", False)])

p = doc.add_paragraph()
run = p.add_run("Same-Day Demand Is Blocked")
run.bold = True

add_body(doc, (
    "From a sample of 700 outbound HRD calls, 37% of all leads needed same-day rentals. "
    "However, the MMR booking time guardrails prevent these customers from completing the flow. "
    "The system enforces a 2-hour buffer from the current time (rounded to 30-minute increments), "
    "plus a 1-hour buffer before the location's closing time. For a branch closing at 4:30pm, "
    "customers on the page after approximately 1:30pm see no available time slots. For a 4:00pm "
    "close, the effective cutoff is approximately 1:00pm."
))

add_body(doc, (
    "These customers — who are often the most urgent and motivated — see no options, "
    "perceive that Hertz cannot help them, and drop out of the flow entirely."
))

p = doc.add_paragraph()
run = p.add_run("Unnecessary Confirmation Step")
run.bold = True

add_body(doc, (
    "After completing Page 2, customers are shown a confirmation pop-up modal that was "
    "originally introduced for a credit card deposit flow. Credit cards are no longer collected, "
    "making this extra step unnecessary friction."
))

p = doc.add_paragraph()
run = p.add_run("Low SMS Click-Through")
run.bold = True

add_body(doc, (
    "Only 32% of customers who receive the MMR SMS click the link. There is currently no "
    "follow-up or reminder mechanism for non-clickers."
))


# ── 4. Intended Behaviour ───────────────────────────────────────────────

add_heading_styled(doc, "4. Intended Behaviour")

add_body(doc, (
    "The MMR flow should accommodate all insurance replacement customer scenarios — not just "
    "those with an existing body shop reservation. Specifically:"
))

add_bullet(doc, [("Customers ", False), ("with or without", True),
                  (" a body shop reservation can complete the MMR flow", False)])
add_bullet(doc, [("Customers who need ", False), ("same-day pickup", True),
                  (" are not blocked — they receive a soft confirmation and are routed to "
                   "HRD for immediate assistance", False)])
add_bullet(doc, [("Pickup location is ", False), ("not limited to body shop", True),
                  (" — customers can select home, work, or other locations", False)])
add_bullet(doc, [("The flow completes in all scenarios", True),
                  (", ensuring every engaged customer gets a commitment touchpoint with Hertz", False)])
add_bullet(doc, [("Non-clickers receive a ", False), ("reminder SMS", True),
                  (" to re-engage them into the funnel", False)])


# ── 5. Current UX (As-Is) ───────────────────────────────────────────────

add_heading_styled(doc, "5. Current User Experience & Flows (As-Is)")

add_body(doc, (
    "The MMR flow currently consists of three steps shown in a progress bar: "
    "(1) My Details and Damaged Vehicle, (2) Start a Reservation, and (3) Confirmation."
))

p = doc.add_paragraph()
run = p.add_run("Page 1 — My Details and Damaged Vehicle")
run.bold = True

add_bullet(doc, [("Customer enters via SMS link with confirmation number and last name", False)])
add_bullet(doc, [("Page displays: phone number, email, insurance company name, "
                   "vehicle make, model, year, and a \"Was your vehicle stolen?\" checkbox", False)])
add_bullet(doc, [("Fields are pre-populated from HLES where available; otherwise the customer must enter them", False)])
add_bullet(doc, [("All vehicle fields are required (no \"unknown\" option)", False)])
add_bullet(doc, [("87% of customers who reach this page complete it", False)])

p = doc.add_paragraph()
run = p.add_run("Page 2 — Start a Reservation")
run.bold = True

add_bullet(doc, [('"Do you need a Hertz rep to pick you up from the body shop?" — Yes / No radio buttons', False)])
add_bullet(doc, [("Hertz Rental Location (dropdown, pre-populated)", False)])
add_bullet(doc, [("Pickup Date (date picker, defaults to next day)", False)])
add_bullet(doc, [("Pickup Time (dropdown, 30-minute increments with guardrails)", False)])
add_bullet(doc, [("Body Shop Location (dropdown with search)", False)])
add_bullet(doc, [("Static car image on the right side — clicking it reveals the branch phone number "
                   "and \"select option 5\" instruction (not intuitive)", False)])
add_bullet(doc, [('"For additional support please reply to text or call 1-800-582-7499" at the bottom', False)])
add_bullet(doc, [("66% of customers who reach Page 2 complete it (34% drop-off)", False)])

p = doc.add_paragraph()
run = p.add_run("Confirmation Pop-Up (Modal)")
run.bold = True

add_bullet(doc, [("After clicking \"Reserve a Rental\", a modal overlay appears summarising: "
                   "pickup from body shop (Yes/No), Hertz Rental Location, Pickup Date, "
                   "Pickup Time, Body Shop Location", False)])
add_bullet(doc, [("Two buttons: \"Make Changes\" and \"Confirm Details\"", False)])
add_bullet(doc, [("Originally introduced for credit card deposit flow — no longer applicable", False)])

p = doc.add_paragraph()
run = p.add_run("Confirmation Page (Page 3)")
run.bold = True

add_bullet(doc, [("Displays: confirmation number, reservation date/time, branch address, "
                   "branch phone number, iCalendar link", False)])
add_bullet(doc, [("Documents required reminder (Driver's License, Credit Card)", False)])
add_bullet(doc, [('"To finalize pick up service, you will need to contact the location '
                   'directly with a minimum of 1 hour advance notice & select option 5."', False)])
add_bullet(doc, [('"For additional support please reply to text or call 1-800-582-7499."', False)])


# ── 6. Target UX (To-Be) ────────────────────────────────────────────────

add_heading_styled(doc, "6. Target User Experience & Flows (To-Be)")

add_body(doc, (
    "The target state restructures Page 2 into a branching flow that accommodates all "
    "customer scenarios, removes the unnecessary confirmation pop-up, and enhances "
    "the final confirmation page."
))

p = doc.add_paragraph()
run = p.add_run("Page 1 — No Changes")
run.bold = True

add_body(doc, (
    "Page 1 remains unchanged. 87% completion rate indicates this page is not a friction point."
))

p = doc.add_paragraph()
run = p.add_run("Page 2 — Restructured Branching Flow")
run.bold = True

add_body(doc, "The proposed decision flow for Page 2 is as follows:")

# Flow diagram as formatted text
doc.add_paragraph("")
flow_text = (
    "Step 1:  \"Do you already have a body shop?\"\n"
    "           ├── No  → Skip to Soft Confirmation (see Path A below)\n"
    "           └── Yes → Continue to Step 2\n"
    "\n"
    "Step 2:  \"Do you need a Hertz rep to pick you up?\"\n"
    "           ├── No  → \"Will come into location\" passed to HLES → Step 4\n"
    "           └── Yes → Continue to Step 3\n"
    "\n"
    "Step 3:  \"Where would you like to be picked up?\"\n"
    "           Options: Repair Shop | Renter Home | Renter Work | Other Location | Dealership\n"
    "           (Mapped to existing HLES pickup location fields)\n"
    "           → Continue to Step 4\n"
    "\n"
    "Step 4:  Select Pickup Date and Pickup Time\n"
    "           ├── Same-day selected + no time slots available → Path B\n"
    "           └── Slots available → Select time → Confirmation Page\n"
    "\n"
    "─── Path A: No Body Shop ───\n"
    "  • Soft confirmation: \"Your reservation request has been submitted.\n"
    "    Once you know your body shop, please update via this link.\n"
    "    A branch representative will contact you.\"\n"
    "  • Bot enters as \"unsuccessful\" in HLES → stays in branch queue\n"
    "  • MMR is marked as complete (no drop-off)\n"
    "\n"
    "─── Path B: Same-Day, No Slots ───\n"
    "  • Soft confirmation: \"Your reservation request has been submitted.\"\n"
    "  • Direct customer to HRD: \"For immediate assistance, please reply\n"
    "    to the text you received or call 1-800-582-7499.\"\n"
    "  • Bot enters as \"unsuccessful\" in HLES → stays in branch queue as urgent\n"
    "  • MMR is marked as complete (no drop-off)\n"
)
p = doc.add_paragraph()
run = p.add_run(flow_text)
run.font.size = Pt(9.5)
run.font.name = "Consolas"

p = doc.add_paragraph()
run = p.add_run("Key design principle: ")
run.bold = True
p.add_run(
    "The branch phone number and HRD contact information must only appear on the "
    "final confirmation page — never mid-flow — to prevent customers from exiting "
    "the MMR funnel to make a phone call."
)

p = doc.add_paragraph()
run = p.add_run("Confirmation Pop-Up — Removed")
run.bold = True

add_body(doc, (
    "The intermediate confirmation modal is removed entirely. Clicking \"Reserve a Rental\" "
    "on Page 2 proceeds directly to the Confirmation Page. The original purpose of this "
    "modal (credit card deposit confirmation) is no longer applicable."
))

p = doc.add_paragraph()
run = p.add_run("Static Car Image — Removed from Page 2")
run.bold = True

add_body(doc, (
    "The static car image on Page 2 is removed. It is not dynamic (same image for every "
    "reservation), may set incorrect vehicle expectations, and its hidden phone number "
    "feature risks pulling customers out of the MMR flow mid-funnel."
))

p = doc.add_paragraph()
run = p.add_run("Confirmation Page — Enhanced")
run.bold = True

add_body(doc, "The confirmation page retains its current information and adds:")

add_bullet(doc, [("For ", False), ("standard flow", True),
                  (": \"Your reservation is confirmed\" with branch phone number, "
                   "HRD contact, and iCalendar link", False)])
add_bullet(doc, [("For ", False), ("same-day / no-slots flow", True),
                  (": \"Your reservation request has been submitted\" (softer language — "
                   "not \"confirmed\") with prominent HRD contact for warm transfer to branch", False)])
add_bullet(doc, [("For ", False), ("no body shop flow", True),
                  (": \"Your reservation request has been submitted. Once you know your "
                   "body shop, please update via this link.\"", False)])


# ── 7. Rules & Constraints ──────────────────────────────────────────────

add_heading_styled(doc, "7. MMR Booking Time Rules & Constraints")

p = doc.add_paragraph()
run = p.add_run("The following booking time guardrails are currently enforced:")
run.bold = True

add_bullet(doc, [("Time slots are in ", False), ("30-minute increments", True)])
add_bullet(doc, [("Minimum ", False), ("2-hour buffer", True),
                  (" from current time, rounded up to the next 30-minute slot", False)])
add_bullet(doc, [("First available slot: ", False), ("1 hour after location open", True)])
add_bullet(doc, [("Last available slot: ", False), ("1 hour before location close", True)])

add_body(doc, (
    "These rules create an effective cutoff for same-day bookings. For example:"
))

add_table(doc,
    headers=["Location Closing Time", "Last Available Slot", "Effective Same-Day Cutoff"],
    rows=[
        ["4:00 PM", "3:00 PM", "~1:00 PM"],
        ["4:30 PM", "3:30 PM", "~1:30 PM"],
        ["5:00 PM", "4:00 PM", "~2:00 PM"],
    ],
    col_widths=[5, 5, 5],
)

doc.add_paragraph("")

add_bold_body(doc, [
    ("Note: ", True),
    ("Exact guardrail parameters to be confirmed with the tech team. The observed behaviour "
     "from the meeting discussion (last available slot of 2:30pm for a 4:30pm close) suggests "
     "the actual constraints may differ slightly from the stated rules.", False),
])

p = doc.add_paragraph()
run = p.add_run("Messaging constraints:")
run.bold = True

add_bullet(doc, [("Standard flow → ", False), ("\"Your reservation is confirmed\"", True)])
add_bullet(doc, [("Same-day / no body shop flow → ", False),
                  ("\"Your reservation request has been submitted\"", True),
                  (" (softer language to avoid confirming a reservation the branch has not yet validated)", False)])
add_bullet(doc, [("Branch phone number shown ", False), ("only on the final page", True),
                  (", never mid-flow", False)])


# ── 8. Functional Requirements ───────────────────────────────────────────

add_heading_styled(doc, "8. Functional Requirements")

add_body(doc, (
    "Below are the main themes across the problem areas. Each theme encompasses "
    "a set of changes to the MMR flow."
))

func_req_rows = [
    ["8.1", "Page 2 Flow Restructure:\nBody Shop & Pickup Branching",
     "Add branching questions to Page 2: \"Do you already have a body shop?\" (Yes/No), "
     "\"Do you need pickup?\" (Yes/No), and \"Where would you like to be picked up?\" "
     "(Repair Shop, Renter Home, Renter Work, Other Location, Dealership). "
     "Each path maps to existing HLES fields. Customers without a body shop skip to a "
     "soft confirmation. Pickup location options expand beyond body-shop-only."],
    ["8.2", "Same-Day Pickup\nException Handling",
     "When a customer selects same-day pickup and no time slots are available due to "
     "booking guardrails, do not block the flow. Complete the MMR with a soft confirmation "
     "(\"request submitted\", not \"confirmed\"). Bot enters as \"unsuccessful\" in HLES so "
     "the reservation stays in the branch queue as needing urgent follow-up. Direct the "
     "customer to HRD (reply to text or call) for warm transfer to the branch."],
    ["8.3", "Remove Confirmation\nPop-Up Modal",
     "Remove the intermediate \"Confirm Reservation Details\" modal that appears after "
     "clicking \"Reserve a Rental\". This modal was introduced for credit card deposit "
     "confirmation, which is no longer collected. Page 2 submission proceeds directly "
     "to the Confirmation Page."],
    ["8.4", "Remove Static Car\nImage from Page 2",
     "Remove the static (non-dynamic) car image from Page 2. The image is the same for "
     "every reservation, may set incorrect vehicle expectations, and contains a hidden "
     "branch phone number that risks pulling customers out of the MMR flow mid-funnel."],
    ["8.5", "Confirmation Page\nEnhancements",
     "Enhance the confirmation page messaging to differentiate between standard confirmation "
     "(\"Your reservation is confirmed\"), same-day/no-slots soft confirmation (\"Your "
     "reservation request has been submitted\" + HRD contact), and no-body-shop soft "
     "confirmation (\"Request submitted, please update once you know your body shop\"). "
     "Branch phone number and HRD contact appear only on this final page."],
    ["8.6", "SMS Reminder for\nNon-Clickers",
     "If no action is taken on the MMR SMS within 6 hours (no click, and no contact from "
     "HRD or branch), send a second automated SMS reminder as an additional touchpoint "
     "to re-engage the customer into the MMR funnel."],
]

table = doc.add_table(rows=1 + len(func_req_rows), cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(["#", "Theme", "Description"]):
    cell = table.rows[0].cells[i]
    cell.text = ""
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(cell, "E8EDF0")

# Rows
for r_idx, row in enumerate(func_req_rows):
    for c_idx, val in enumerate(row):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx <= 1:
            run.bold = True

# Column widths
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(4.5)
    row.cells[2].width = Cm(10)


# ── 9. Non-Functional Requirements ──────────────────────────────────────

add_heading_styled(doc, "9. Non-Functional Requirements")

add_bullet(doc, [("No perceptible delay when branching questions are displayed "
                   "(client-side logic, no additional API calls for branching)", False)])
add_bullet(doc, [("Soft confirmation paths must still write to HLES correctly "
                   "(\"unsuccessful\" status with appropriate pickup notes)", False)])
add_bullet(doc, [("SMS reminder system must respect opt-out / do-not-contact preferences", False)])
add_bullet(doc, [("All flow changes must be backward-compatible with existing HLES "
                   "reservation processing", False)])
add_bullet(doc, [("Pickup location options must map to existing HLES fields "
                   "(Repair Shop, Renter Home, Renter Work, Other Location, Dealership, "
                   "Will Come Into Location)", False)])


# ── 10. Success Metrics & Tracking Plan ──────────────────────────────────

add_heading_styled(doc, "10. Success Metrics & Tracking Plan")

p = doc.add_paragraph()
run = p.add_run("Uplift Targets")
run.bold = True

add_table(doc,
    headers=["Metric", "Current", "Target (Low)", "Target (High)"],
    rows=[
        ["MMR SMS Click-Through Rate", "32.2%", "40%", "50%"],
        ["MMR Page 2 Completion Rate", "66%", "82.5%", "89.1%"],
        ["Conversion Uplift (Page 1 initiative)", "Baseline", "+0.61pp", "+1.40pp"],
        ["Conversion Uplift (Page 2 initiative)", "Baseline", "+0.63pp", "+0.89pp"],
        ["Combined Conversion Uplift", "Baseline", "+1.24pp", "+2.29pp"],
        ["EBITDA Uplift (MMR combined)", "—", "$1.2M", "$2.3M"],
        ["Revenue Uplift (MMR combined)", "—", "$3.3M", "$6.2M"],
    ],
    col_widths=[5.5, 3, 3.5, 3.5],
)

doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("Weekly Tracking Plan")
run.bold = True

add_body(doc, "The following metrics will be tracked on a weekly basis using Anupam's data:")

add_bullet(doc, [("MMR click-through rate", True), (" — % of SMS recipients who click the MMR link", False)])
add_bullet(doc, [("MMR Page 1 completion rate", True), (" — % of clickers who complete Page 1", False)])
add_bullet(doc, [("MMR Page 2 completion rate", True), (" — % of Page 1 completers who complete Page 2", False)])
add_bullet(doc, [("Conversion rate with MMR", True),
                  (" — % conversion for customers who completed MMR vs. those who did not", False)])
add_bullet(doc, [("Net % of reservations with MMR completions", True),
                  (" — total share of reservations that went through completed MMR flow", False)])


# ── 11. Out of Scope ─────────────────────────────────────────────────────

add_heading_styled(doc, "11. Out of Scope")

add_bullet(doc, [("Changes to Page 1 (My Details and Damaged Vehicle)", False)])
add_bullet(doc, [("Changes to the SMS message content or format", False)])
add_bullet(doc, [("Changes to HLES backend reservation processing logic", False)])
add_bullet(doc, [("A/B testing infrastructure (not currently available for MMR)", False)])
add_bullet(doc, [("Drivable / non-drivable vehicle status capture (no structured HLES field exists "
                   "to store this at reservation confirmation time)", False)])
add_bullet(doc, [("HRD operational process changes (warm transfer training, call transfer protocols)", False)])
add_bullet(doc, [("Branch operational process changes (phone pickup protocols)", False)])


# ── Final Note ───────────────────────────────────────────────────────────

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Final Note")
run.bold = True

add_bold_body(doc, [
    ("This PRD positions MMR optimisation as a ", False),
    ("high-impact, low-friction initiative", True),
    (" to unlock conversion lift from customers who are already engaged. By restructuring "
     "Page 2 to accommodate all customer scenarios — not just the body-shop-ready happy "
     "path — and by introducing a reminder mechanism for non-clickers, Hertz can capture "
     "incremental revenue while improving the customer experience for insurance replacement "
     "renters.", False),
])


# ── Save ─────────────────────────────────────────────────────────────────

output_dir = "/Users/dansia/Documents/HertzDataAnalysis/docs/MMR initiative"
output_path = os.path.join(output_dir, "PRD - MMR Flow Optimisation.docx")
doc.save(output_path)
print(f"PRD saved to: {output_path}")
